from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path('/workspace/scratch/22e4ba00bbb2')
SNAPSHOT = ROOT / 'atc-main-snapshot'
OUT_DIR = ROOT / 'deliverables'
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT = OUT_DIR / 'ATC_Master_Handoff_2026-08-26.docx'

DOC_SKILL = Path('/root/.codex/skills/builtins/documents')
sys.path.insert(0, str(DOC_SKILL / 'scripts'))
from table_geometry import apply_table_geometry  # noqa: E402


# compact_reference_guide preset, plus named cover/status overrides.
BLUE = '2E74B5'
DARK_BLUE = '1F4D78'
INK = '0B2545'
MUTED = '5E6B75'
LIGHT_BLUE = 'E8EEF5'
LIGHT_GRAY = 'F2F4F7'
CALLOUT = 'F4F6F9'
GREEN = '147D64'  # named brand override: Bumba/Airspace status accent
DARK_GREEN = '0B3B2E'
WHITE = 'FFFFFF'
RED = '9B1C1C'
GOLD = '7A5A00'
TABLE_WIDTH = 9360
TABLE_INDENT = 120
CELL_MARGINS = {'top': 80, 'bottom': 80, 'start': 120, 'end': 120}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_text_color(cell, color: str) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement('w:tblHeader')
    header.set(qn('w:val'), 'true')
    tr_pr.append(header)


def set_run_font(run, *, name='Calibri', size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_hyperlink(paragraph, text: str, url: str, color=BLUE):
    part = paragraph.part
    rel_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), rel_id)
    run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')
    r_color = OxmlElement('w:color')
    r_color.set(qn('w:val'), color)
    r_pr.append(r_color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    r_pr.append(underline)
    r_fonts = OxmlElement('w:rFonts')
    r_fonts.set(qn('w:ascii'), 'Calibri')
    r_fonts.set(qn('w:hAnsi'), 'Calibri')
    r_pr.append(r_fonts)
    run.append(r_pr)
    text_node = OxmlElement('w:t')
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run('Sayfa ')
    set_run_font(run, size=8.5, color=MUTED)
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_numbering(doc: Document):
    numbering = doc.part.numbering_part.element

    def create_abstract(abstract_id: int, fmt: str, text: str):
        abstract = OxmlElement('w:abstractNum')
        abstract.set(qn('w:abstractNumId'), str(abstract_id))
        multi = OxmlElement('w:multiLevelType')
        multi.set(qn('w:val'), 'singleLevel')
        abstract.append(multi)
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), '0')
        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)
        num_fmt = OxmlElement('w:numFmt')
        num_fmt.set(qn('w:val'), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement('w:lvlText')
        lvl_text.set(qn('w:val'), text)
        lvl.append(lvl_text)
        jc = OxmlElement('w:lvlJc')
        jc.set(qn('w:val'), 'left')
        lvl.append(jc)
        p_pr = OxmlElement('w:pPr')
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'num')
        tab.set(qn('w:pos'), '540')
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '540')
        ind.set(qn('w:hanging'), '270')
        p_pr.append(ind)
        lvl.append(p_pr)
        abstract.append(lvl)
        numbering.append(abstract)

    def create_num(num_id: int, abstract_id: int):
        num = OxmlElement('w:num')
        num.set(qn('w:numId'), str(num_id))
        abstract_num_id = OxmlElement('w:abstractNumId')
        abstract_num_id.set(qn('w:val'), str(abstract_id))
        num.append(abstract_num_id)
        numbering.append(num)

    create_abstract(90, 'bullet', '•')
    create_num(90, 90)
    create_abstract(91, 'decimal', '%1.')
    create_num(91, 91)
    return {'bullet': 90, 'number': 91, 'number_abstract': 91, 'next_num_id': 92}


def reset_number_sequence(doc: Document, num_ids) -> None:
    """Start a fresh real Word-numbered list at 1."""
    num_id = num_ids['next_num_id']
    num_ids['next_num_id'] += 1
    num_ids['number'] = num_id
    numbering = doc.part.numbering_part.element
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(num_id))
    abstract_num_id = OxmlElement('w:abstractNumId')
    abstract_num_id.set(qn('w:val'), str(num_ids['number_abstract']))
    num.append(abstract_num_id)
    level_override = OxmlElement('w:lvlOverride')
    level_override.set(qn('w:ilvl'), '0')
    start_override = OxmlElement('w:startOverride')
    start_override.set(qn('w:val'), '1')
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)


def apply_num(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn('w:numPr'))
    if num_pr is None:
        num_pr = OxmlElement('w:numPr')
        p_pr.append(num_pr)
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    num = OxmlElement('w:numId')
    num.set(qn('w:val'), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def add_bullet(doc, text: str, num_ids, *, bold_prefix: str | None = None):
    paragraph = doc.add_paragraph()
    apply_num(paragraph, num_ids['bullet'])
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    set_paragraph_spacing(paragraph, after=4, line=1.25)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_run_font(run, bold=True, color=INK)
        run = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(run)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_numbered(doc, text: str, num_ids):
    paragraph = doc.add_paragraph()
    apply_num(paragraph, num_ids['number'])
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    set_paragraph_spacing(paragraph, after=4, line=1.25)
    set_run_font(paragraph.add_run(text))
    return paragraph


def add_body(doc, text: str, *, bold_lead: str | None = None, italic=False):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph)
    if bold_lead and text.startswith(bold_lead):
        set_run_font(paragraph.add_run(bold_lead), bold=True, color=INK)
        set_run_font(paragraph.add_run(text[len(bold_lead):]), italic=italic)
    else:
        set_run_font(paragraph.add_run(text), italic=italic)
    return paragraph


def add_callout(doc, label: str, text: str, *, tone='info'):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=5, after=9, line=1.2)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), CALLOUT if tone == 'info' else ('FFF4DE' if tone == 'warning' else 'FDECEC'))
    p_pr.append(shd)
    borders = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), GREEN if tone == 'info' else (GOLD if tone == 'warning' else RED))
    borders.append(left)
    p_pr.append(borders)
    set_run_font(paragraph.add_run(f'{label}: '), bold=True, color=INK)
    set_run_font(paragraph.add_run(text))
    return paragraph


def set_table_borders(table, color='CBD4DC', size='4'):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = qn(f'w:{edge}')
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f'w:{edge}')
            borders.append(node)
        node.set(qn('w:val'), 'single')
        node.set(qn('w:sz'), size)
        node.set(qn('w:space'), '0')
        node.set(qn('w:color'), color)


def add_table(doc, headers, rows, widths, *, font_size=9.2, alignments=None, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for index, value in enumerate(headers):
        cell = hdr.cells[index]
        cell.text = str(value)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, header_fill)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = (alignments[index] if alignments else WD_ALIGN_PARAGRAPH.LEFT)
        set_paragraph_spacing(paragraph, after=0, line=1.05)
        for run in paragraph.runs:
            set_run_font(run, size=font_size, color=INK, bold=True)
    for row_data in rows:
        row = table.add_row()
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (alignments[index] if alignments else WD_ALIGN_PARAGRAPH.LEFT)
            set_paragraph_spacing(paragraph, after=0, line=1.08)
            for run in paragraph.runs:
                set_run_font(run, size=font_size, color='1F2933')
    set_table_borders(table)
    apply_table_geometry(table, widths, table_width_dxa=TABLE_WIDTH, indent_dxa=TABLE_INDENT, cell_margins_dxa=CELL_MARGINS)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    return table


def parse_airports():
    text = (SNAPSHOT / 'docs/AIRPORT_DATA.md').read_text(encoding='utf-8')
    rows = []
    for line in text.splitlines():
        match = re.match(r'\|\s*(\d+)\s*\|\s*([A-Z0-9]+)\s*\|\s*([A-Z0-9]+)\s*\|\s*([\d,]+)\s*\|\s*(\d+)\s*\|', line)
        if match:
            rank, iata, icao, passengers, runways = match.groups()
            rows.append((int(rank), iata, icao, passengers, int(runways)))
    if len(rows) != 50:
        raise RuntimeError(f'Expected 50 airport rows, found {len(rows)}')
    return rows


def parse_published_procedures():
    text = (SNAPSHOT / 'docs/PUBLISHED_PROCEDURES.md').read_text(encoding='utf-8')
    rows = []
    in_table = False
    for line in text.splitlines():
        if line.startswith('| Airport | Runtime procedures'):
            in_table = True
            continue
        if in_table and line.startswith('| ---'):
            continue
        if in_table and not line.startswith('|'):
            break
        if in_table:
            parts = [part.strip() for part in line.strip().strip('|').split('|')]
            if len(parts) == 3:
                rows.append(tuple(parts))
    if len(rows) != 29:
        raise RuntimeError(f'Expected 29 procedure rows, found {len(rows)}')
    return rows


def parse_achievements():
    text = (SNAPSHOT / 'src/engine/progression.ts').read_text(encoding='utf-8')
    block = text.split('export const ACHIEVEMENTS', 1)[1].split('];', 1)[0]
    rows = re.findall(r"\{ id: '([^']+)', label: '([^']+)', description: '([^']+)'", block)
    if len(rows) != 52:
        raise RuntimeError(f'Expected 52 achievements, found {len(rows)}')
    return rows


AIRPORTS = parse_airports()
PROCEDURES = parse_published_procedures()
ACHIEVEMENTS = parse_achievements()
PUBLISHED_ICAO = {row[0] for row in PROCEDURES}


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles['Title']
    title.font.name = 'Calibri'
    title._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    title._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(DARK_GREEN)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)

    subtitle = styles['Subtitle']
    subtitle.font.name = 'Calibri'
    subtitle._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    subtitle._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_after = Pt(18)

    specs = {
        'Heading 1': (16, BLUE, 18, 10),
        'Heading 2': (13, BLUE, 14, 7),
        'Heading 3': (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in specs.items():
        style = styles[name]
        style.font.name = 'Calibri'
        style._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
        style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_header_footer(section):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(paragraph, after=0, line=1)
    set_run_font(paragraph.add_run('AIRSPACE CONTROL  ·  MASTER HANDOFF'), size=8.5, color=MUTED, bold=True)
    footer = section.footer
    add_page_field(footer.paragraphs[0])


doc = Document()
configure_styles(doc)
for section in doc.sections:
    configure_section(section)
    configure_header_footer(section)
num_ids = add_numbering(doc)

doc.core_properties.title = 'Airspace Control - Master Handoff'
doc.core_properties.subject = 'Product, research, data, technical architecture, deployment and next-step handoff'
doc.core_properties.author = 'Bumba Games / Buğra Bilim'
doc.core_properties.keywords = 'ATC, Airspace Control, Bumba Games, GitHub, Vercel, handoff'
doc.core_properties.comments = 'Snapshot verified 2026-08-26. Game-only; not for navigation.'


# Cover page: editorial_cover pattern with a compact status strip.
spacer = doc.add_paragraph()
spacer.paragraph_format.space_after = Pt(62)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p, after=14)
set_run_font(p.add_run('BUMBA GAMES  ·  PRODUCT & TECHNICAL HANDOFF'), size=10, color=GREEN, bold=True)

p = doc.add_paragraph(style='Title')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('AIRSPACE CONTROL')
p = doc.add_paragraph(style='Subtitle')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Master project file for continuing development in a new conversation')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p, after=24)
set_run_font(p.add_run('Snapshot: 26 August 2026  ·  Canonical branch: GitHub main'), size=10.5, color=MUTED, bold=True)

cover_rows = [
    ('50', 'PLAYABLE AIRPORTS', '29', 'PUBLISHED ROUTE AIRPORTS', '137', 'PASSING TESTS'),
]
table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'
set_repeat_table_header(table.rows[0])
for idx, value in enumerate(cover_rows[0]):
    cell = table.rows[0].cells[idx]
    cell.text = value
    set_cell_shading(cell, DARK_GREEN if idx % 2 == 0 else LIGHT_BLUE)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, after=0, line=1)
    for run in paragraph.runs:
        set_run_font(run, size=18 if idx % 2 == 0 else 8, color=WHITE if idx % 2 == 0 else INK, bold=True)
apply_table_geometry(table, [800, 2320, 800, 2320, 800, 2320], table_width_dxa=9360, indent_dxa=100, cell_margins_dxa={'top': 150, 'bottom': 150, 'start': 100, 'end': 100})
set_table_borders(table, color='8BA99C', size='5')

p = doc.add_paragraph()
set_paragraph_spacing(p, before=28, after=5)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run('Current production'), size=9, color=MUTED, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p, after=2)
add_hyperlink(p, 'Canlı ürünü aç', 'https://atc-tr.vercel.app', color=GREEN)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p, after=1)
add_hyperlink(p, 'github.com/bugrabilim/atc', 'https://github.com/bugrabilim/atc')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p, before=28, after=0)
set_run_font(p.add_run('GAME ONLY  ·  NOT FOR NAVIGATION OR OPERATIONAL ATC'), size=8.5, color=RED, bold=True)

doc.add_page_break()


doc.add_heading('0. Yeni sohbette nasıl devam edilir?', level=1)
add_callout(doc, 'YENİ SOHBET BAŞLANGIÇ METNİ',
            'Bu dosyayı yeni sohbete yükle ve şunu yaz: “Bu dosya Airspace Control projesinin ana devir belgesidir. Önce GitHub main ve canlı Vercel durumunu yeniden doğrula. Eski yerel dalı kaynak kabul etme ve hiçbir yerel değişikliği silme. Depoya veya Vercel üretimine göndermeden önce benden onay al. Ardından belgedeki P0 işlerden sırayla devam et.”')
add_body(doc, 'Bu belge sohbet geçmişine ihtiyaç duymadan proje kararlarını, mevcut uygulamayı, doğrulanmış teknik durumu ve yapılacak işleri taşımak için hazırlanmıştır. Yeni sohbette güncel kaynak olarak önce GitHub main kontrol edilmeli; bu belgedeki sayısal durum daha yeni bir commit varsa yeniden hesaplanmalıdır.')

doc.add_heading('Belge haritası', level=2)
for item in [
    '1. Yönetici özeti ve ürün hedefi',
    '2. Sabit ürün kararları ve sınırlar',
    '3. GitHub, Vercel ve çalışma alanı durumu',
    '4. Oyuncu deneyimi ve mevcut ekranlar',
    '5. Simülasyon, trafik, emniyet ve komut sistemleri',
    '6. İlerleme, Academy, kariyer, Daily Radar ve başarımlar',
    '7. Havalimanı, pist, coğrafya ve yayımlanmış prosedür verisi',
    '8. Teknik mimari, kayıt, PWA ve mobil paketleme',
    '9. FlatOut ATC / Endless ATC araştırma sonucu ve özgünlük sınırı',
    '10. Doğrulama sonuçları, bilinen açıklar ve riskler',
    '11. Öncelikli uygulama planı ve tamamlanma ölçütleri',
    'Ekler: 50 havalimanı, 29 prosedür paketi ve 52 başarım',
]:
    add_bullet(doc, item, num_ids)


doc.add_heading('1. Yönetici özeti ve ürün hedefi', level=1)
add_callout(doc, 'MEVCUT GERÇEK DURUM', 'Airspace Control çalışan, canlıya alınmış ve testli bir 2D radar ATC oyun temelidir. Ancak FlatOut ATC ürün derinliğine, fiziksel mobil cihaz kalitesine ve 50 havalimanının tamamında yayımlanmış prosedür kapsamına henüz ulaşmamıştır.')

summary_rows = [
    ('Ürün', 'Airspace Control; Bumba Games markası altında özgün ATC radar strateji oyunu.'),
    ('Platform', 'Web ve PWA çalışıyor. Capacitor yapılandırması var; mağaza için imzalı AAB/IPA henüz üretilmedi.'),
    ('Canlı adres', 'atc-tr.vercel.app ana landing; atc-tr-play.vercel.app doğrudan oyun görünümünü açan alias.'),
    ('Depo', 'github.com/bugrabilim/atc; kaynak gerçekliği origin/main.'),
    ('Çekirdek', 'Seed’li random trafik, adaptif skill, 2D radar, heading/irtifa/hız, DCT/HOLD, STAR/SID, ILS/LOC, wake, ayırma, CPA, debrief.'),
    ('İçerik', '50 havalimanı, 4 zorluk, 10 Academy dersi, 7 bölümlük First Watch, Daily Radar, 30 vardiyalık logbook, 52 başarım.'),
    ('Veri kapsamı', '50/50 pist kataloğu; 29/50 havalimanında yayımlanmış runtime rotası (%58); 21 havalimanında üretilmiş vektör.'),
    ('Bulut', 'Supabase, hesap, cloud sync ve online leaderboard kullanıcı kararıyla ertelendi.'),
    ('Trafik', 'Yapay zekâ kullanılmıyor; deterministik/seed’li random trafik yeterli kabul edildi.'),
    ('Durum', 'Üretim READY; son 24 saatte Vercel runtime hatası yok; 137 test geçiyor.'),
]
add_table(doc, ['Alan', 'Doğrulanmış durum'], summary_rows, [1900, 7460], font_size=9.5)

doc.add_heading('1.1 Hedef ürün tanımı', level=2)
add_body(doc, 'Hedef, FlatOut ATC ve Endless ATC’nin oyuncuya sunduğu karar yoğunluğu, erişilebilirlik, eğitim, havaalanı karakteri ve uzun süreli ilerleme derinliğine ulaşan; fakat bağımsız kod, özgün arayüz, özgün hikâye, özgün metin/ses/görsel ve belgelenmiş veri kaynakları kullanan bir oyundur.')
add_body(doc, '“İsim ve renkleri değiştirerek birebir kopya” güvenli bir hedef değildir. Apple Copycats kuralı ve fikrî mülkiyet riski nedeniyle işlevsel kapsam eşleşebilir; ekran yerleşimi, metinler, ikonlar, görseller, sesler, görev yazımı ve kod kopyalanamaz.')

doc.add_heading('1.2 Temel ürün vaadi', level=2)
for text in [
    'Yeni oyuncu iki dakika içinde ilk uçağını seçip güvenli bir yaklaşma başlatabilmeli.',
    'Deneyimli oyuncu heading, irtifa, hız, DCT/HOLD ve yaklaşma komutlarıyla yoğun bir sektörü hızlı yönetebilmeli.',
    'Her havalimanı yalnızca farklı pist çizimi değil, kendine özgü akış, kapasite, olay ve trafik karakteri sunmalı.',
    'Puan, başarımlar, Academy, Daily Radar, logbook ve kariyer bölümleri tekrar oynama nedeni oluşturmalı.',
    'Mobil ekran radarı küçültmemeli; komutlar büyük dokunma alanlarıyla bağlamsal bir alt panelde açılmalı.',
]:
    add_bullet(doc, text, num_ids)


doc.add_heading('2. Sabit ürün kararları ve sınırlar', level=1)
decisions = [
    ('Marka', 'Ürün adı Airspace Control; stüdyo/alt marka Bumba Games.'),
    ('Ana hedef', 'FlatOut ATC düzeyinde işlevsel ürün derinliği ve Endless ATC düzeyinde akıcı radar kontrolü; görsel veya metinsel kopya değil.'),
    ('Platform', 'Mobil ve web aynı oyun motorunu kullanacak. PWA hemen kullanılabilir; native mağaza paketi daha sonra.'),
    ('Giriş', 'Şimdilik şifresiz, hesapsız ve local-first.'),
    ('Trafik', 'Yapay zekâ gerekmez; seed’li random trafik tercih edilir.'),
    ('Zorluk', 'Başlangıç, Normal, İleri ve Uzman olmak üzere dört mod; seçenekler moda göre artar.'),
    ('Kariyer', 'Mini Metro benzeri ardışık havalimanı açma; önceki meydandaki skor sonraki meydanı açar.'),
    ('Katalog', '50 yoğun havalimanı; İstanbul ürün kararıyla birinci, kalanlar 2025 yolcu sırasına göre.'),
    ('Ses', 'Varsayılan açık; çağrı kodları NATO alfabesi ve rakamlar tek tek okunur. Tarayıcı kullanıcı etkileşimi olmadan sesi başlatamaz.'),
    ('Dil', 'Landing page İngilizce varsayılan; oyun içi kontrol ve yardım metinleri ağırlıklı Türkçe.'),
    ('Veri', 'Pist ve prosedürler kamuya açık, doğrulanabilir kaynaklardan; lisans/provenans kaydı tutulur. Oyun verisi seyrüsefer için değildir.'),
    ('Bulut', 'Supabase ve çevrimiçi leaderboard şimdilik atlanmıştır.'),
    ('Mağaza', 'Google Play/App Store yayını ve imzalı paketler sonraki aşamadır.'),
    ('Değişiklik güvenliği', 'GitHub main veya Vercel production’a gönderme işlemi açık onay olmadan yapılmamalıdır.'),
]
add_table(doc, ['Karar alanı', 'Geçerli karar'], decisions, [1900, 7460], font_size=9.2)


doc.add_heading('3. GitHub, Vercel ve çalışma alanı durumu', level=1)
repo_rows = [
    ('GitHub repo', 'https://github.com/bugrabilim/atc'),
    ('Canonical branch', 'main / origin/main'),
    ('Son commit', 'b95ec92ab64e6098f66f13b3ff72c43810aafe81'),
    ('Commit mesajı', 'feat: add all Doha RNP STARs'),
    ('Commit zamanı', '24 Ağustos 2026 18:15 +03:00'),
    ('Vercel proje', 'atc · prj_vY0tkOUcAiwQfM2OTlgLPugvroKi'),
    ('Vercel takım', 'team_1CGjxTCQLZ149PrgMVSxuVP7'),
    ('Son deployment', 'dpl_4FVj58VvhSPxFVE7Nrw4S1RQC59D · production · READY'),
    ('Primary domain', 'https://atc-tr.vercel.app'),
    ('Direct-play alias', 'https://atc-tr-play.vercel.app'),
    ('Runtime health', '26 Ağustos 2026 kontrolünde son 24 saatte runtime hata kümesi yok.'),
]
add_table(doc, ['Alan', 'Değer'], repo_rows, [2000, 7360], font_size=9.3)

doc.add_heading('3.1 Yerel dal uyarısı', level=2)
add_callout(doc, 'KRİTİK', 'Mevcut yerel çalışma kopyası feat/academy-mobile-core dalında ve değişiklik içeriyor; GitHub main’den geride. Yeni sohbette bu dal resetlenmemeli, checkout ile ezilmemeli veya körlemesine merge edilmemeli. Önce değişiklikler karşılaştırılmalı; geliştirme için temiz bir origin/main çalışma kopyası oluşturulmalıdır.', tone='warning')
add_body(doc, 'Yerelde değiştirilmiş görünen dosyalar: airportOperations.test.ts, airportOperations.ts, navigation.test.ts, navigation.ts, scenario.test.ts, scenario.ts, trafficDirector.test.ts, trafficDirector.ts, types.ts, CommandPanel.tsx. Ayrıca docs/PUBLISHED_PROCEDURES.md untracked görünmektedir. Bunlar atılmadan önce main ile içerik karşılaştırması yapılmalıdır.')

doc.add_heading('3.2 Son geliştirme dalgası', level=2)
recent_rows = [
    ('47cd2d2', 'DFW, ORD ve DEN için FAA CIFP gelişleri'),
    ('94a62e0', 'Kalan ABD havalimanları için FAA arrivals'),
    ('fc56fcf → 41a6a0f', 'Delhi, Incheon ve Dubai yayımlanmış STAR paketleri'),
    ('3bb7a1e → 47143d4', 'Paris, Singapore ve Amsterdam paketleri'),
    ('5fa28d7 → f18f11a', 'Madrid ve Kuala Lumpur paketleri'),
    ('b9e8025 → 13cf139', 'Bangkok ve Hong Kong paketleri'),
    ('ee4329b', 'Barcelona 39 RNAV1 STAR'),
    ('f1cb0a1', 'Mumbai 52 RNAV1 STAR'),
    ('b95ec92', 'Doha 41 RNP STAR'),
]
add_table(doc, ['Commit/dalga', 'İçerik'], recent_rows, [1800, 7560], font_size=9.2)


doc.add_heading('4. Oyuncu deneyimi ve mevcut ekranlar', level=1)
doc.add_heading('4.1 Landing page', level=2)
for text in [
    'İngilizce varsayılan; sticky header ve gerçek footer mevcut.',
    'Bumba Games markası, “The sky is yours to control” ana mesajı ve Play Now / Resume Shift / Start Academy eylemleri bulunuyor.',
    'Şifresiz kullanım, web + mobile ve free-to-start mesajları açıkça gösteriliyor.',
    'Career, Daily Radar, Academy ve 50 havalimanlık arama/filtreleme bölümleri aynı sayfada.',
    'Mobilde tek kolon, masaüstünde çok kolon; landing kendi dikey scroll konteynerini kullanıyor.',
]:
    add_bullet(doc, text, num_ids)

doc.add_heading('4.2 Oyun ekranı', level=2)
for text in [
    'Radar-first çalışma alanı: Canvas 2D radar ana yüzeydir.',
    'Üst bar puan, iniş, trafik, rüzgâr ve saat bilgisini taşır.',
    'Operasyon barında pist akışı, zorluk, rüzgâr, görüş ve QNH bulunur.',
    'Uçak radar hedefinden veya uçuş listesinden seçilebilir.',
    'Masaüstünde komut konsolu; telefonda seçili uçakla açılan sabit alt sheet bulunur.',
    'Heading ±10°/±30°, irtifa ±1.000 ft ve hız ±20 kt dokunmatik kontrolleri vardır.',
    'Pan, zoom, lock, hareketli etiket, predicted path ve mesafe ölçümü radar araçları arasındadır.',
    'Görev/koç bilgileri radar alanını kaplamamak için katlanabilir/bağlamsal tasarlanmıştır.',
]:
    add_bullet(doc, text, num_ids)

doc.add_heading('4.3 Temel oynanış döngüsü', level=2)
reset_number_sequence(doc, num_ids)
for step in [
    'Oyuncu landing page’den Academy, Daily Radar, First Watch veya serbest havalimanı vardiyası seçer.',
    'Mod ve pist akışı seçilir; senaryo seed’i ile random geliş/kalkış trafiği oluşturulur.',
    'Oyuncu uçağı seçer ve heading, irtifa, hız, route veya yaklaşma talimatı verir.',
    'Pilot readback kısa deterministik gecikmeden sonra komutu uygular.',
    'Oyuncu 3 NM/1.000 ft ayırmayı, wake aralığını, pist kapasitesini ve öncelikli trafiği korur.',
    'ILS; ARMED → LOC → GS → TOWER durumlarıyla ilerler ve kule devri otomatik tamamlanır.',
    'İniş ve handoff puan kazandırır; hatalar puan ve canlı skill değerini azaltır.',
    'Skill yükselince hedef trafik artar; hata sonrası yük doğal olarak azalır.',
    'Vardiya bitince debrief, başarımlar, en iyi skor, kariyer ilerlemesi ve logbook güncellenir.',
]:
    add_numbered(doc, step, num_ids)

doc.add_heading('4.4 Komut söz dizimi', level=2)
command_rows = [
    ('Heading', 'AR101 H090 / HDG 090; opsiyonel L veya R yönü'),
    ('İrtifa', 'AR101 A30 / ALT 3000 / FL060'),
    ('Hız', 'AR101 S220 / SPD 220'),
    ('Yaklaşma', 'AR101 I34L / ILS 34L'),
    ('Localizer only', 'AR101 L34L / LOC 34L / LLZ 34L'),
    ('Direct', 'AR101 DCT FIX'),
    ('Hold', 'AR101 HOLD FIX'),
    ('Prosedür', 'AR101 STAR RIXEN1W / SID VICEN1S'),
    ('Handoff', 'AR101 HANDOFF / HOF'),
    ('Go-around', 'AR101 GA'),
    ('Normal speed', 'AR101 RN / RESUME NORMAL'),
    ('Expedite', 'AR101 X / EXPEDITE'),
    ('Birleşik komut', 'AR101 H090 A30 S180 I34L'),
    ('Callsign kısaltma', 'Aktif trafikte tek eşleşme varsa AR gibi ön ek çözümlenir; Tab tamamlar.'),
]
add_table(doc, ['İşlev', 'Örnek / davranış'], command_rows, [1900, 7460], font_size=9.2)
add_callout(doc, 'LAND KOMUTU', 'Parser tarihsel alias’ı tanıyabilse de standart yaklaşma kontrol modunda LAND kullanılmaz. Localizer ve glideslope established olduğunda sistem uçağı otomatik olarak tower durumuna geçirir.')

doc.add_heading('4.5 Ses ve radyo', level=2)
for text in [
    'Audio başlangıç durumu açıktır; kullanıcı arayüzünde “SES AÇIK” görünür.',
    'Tarayıcı güvenliği nedeniyle AudioContext ve speech synthesis ilk kullanıcı etkileşiminden sonra etkinleşebilir; sayfa açılır açılmaz otomatik ses garanti edilemez.',
    'NATO alfabesi kullanılır: CF101 → “Charlie Foxtrot one zero one”. Dokuz “niner” olarak okunur.',
    'Komut, readback, handoff, landing, warning ve alert için offline sentezlenmiş ses ipuçları vardır.',
    'TTS öncelikle İngilizce sesleri, yoksa Türkçe sesleri kullanır; çağrı kodları ve pistler ATC tarzında dönüştürülür.',
    'Aynı uçağa birleşik komutlar tek radyo iletimi olarak kuyruğa alınır; talimatlar ayrı ayrı uygulanabilir kalır.',
]:
    add_bullet(doc, text, num_ids)


doc.add_heading('5. Simülasyon, trafik, emniyet ve olay sistemleri', level=1)
systems_rows = [
    ('Clock', '0,05 saniyelik sabit simülasyon adımı; UI çiziminden ayrıdır.'),
    ('Fizik', 'Bank/roll, IAS/ground speed, rüzgâr vektörü, hız tabanlı dönüş yarıçapı, dikey hız ve uçak performans sınıfları.'),
    ('Yaklaşma', '18 NM’ye kadar yakalama; açı/koridor/glideslope koşulları; LOC, GS, tower ve go-around.'),
    ('Navigation', 'Route, direct ve hold modları; fix geçişi; published fix irtifa/hız kısıtları.'),
    ('Ayırma', '5 NM/1.500 ft uyarı; 3 NM/1.000 ft kayıp; 120 sn CPA öngörüsü.'),
    ('İstisnalar', 'Bağımsız paralel yaklaşmalar, diverging departures ve go-around grace.'),
    ('Wake', 'A–F altı kategori; lider/takipçi matrisi 3–8 NM; kalkış wake süreleri 60–180 sn.'),
    ('Pist', '45 sn turnaround; pist kapanması/yeniden açılması ve aktif flow kapasitesi.'),
    ('Trafik', 'Seed’li callsign/fleet seçimi; airport-specific cadence; heavy cadence; published route tercihi.'),
    ('Skill', 'Başarılar skill’i yükseltir; ayırma, go-around, kaçan handoff/arrival, expired priority ve wake ihlali düşürür.'),
    ('Olaylar', 'Demand pulse, operasyonel flow change/recovery, expert runway inspection, terminal metering ve priority traffic.'),
    ('Öncelik', 'Medical veya minimum fuel; cömert fakat ölçülen response window.'),
]
add_table(doc, ['Alt sistem', 'Mevcut davranış'], systems_rows, [1750, 7610], font_size=9.1)

doc.add_heading('5.1 Skill ve trafik geri beslemesi', level=2)
add_body(doc, 'Başlangıç genel skill değeri 3,5’tir; minimum 2, maksimum 30’dur. Tower handoff +0,55, departure handoff +0,12 sağlar. Ayırma kaybı −1,35; go-around −0,45; missed handoff −0,50; unmanaged arrival −0,65; expired priority −0,55; wake ihlali −0,40 etkiler.')
add_body(doc, 'Hedef aktif uçak sayısı kabaca skill’in yuvarlanmış değeridir ve 3–24 arasında tutulur; pist akışı kapasite düzeltmesi ekler. Spawn aralığı skill yükseldikçe 21 saniyeden 6,5 saniyeye doğru sıklaşır ve mod sınırları uygulanır.')

doc.add_heading('5.2 Skor', level=2)
add_body(doc, 'Ödül = iniş ×125 + handoff ×55 + peak skill ×15. Cezalar: separation loss ×260, wake violation ×100, go-around ×45, missed handoff ×85, unmanaged arrival ×120, expired priority ×170. Son skor sıfırın altına inmez.')

doc.add_heading('5.3 Prosedür kullanan trafik davranışı', level=2)
add_callout(doc, 'TASARIM GERİLİMİ', '29 havalimanında random gelişler uyumlu bir yayımlanmış STAR/stack rotasını tercih eder ve oyuncu vector, DCT, HOLD veya ILS verene kadar rotayı izler. Endless tarzı saf vector-first oynanış ile FlatOut tarzı procedure-assisted oynanış arasında bu davranış yeni sohbette bilinçli olarak korunmalı veya değiştirilmelidir.', tone='warning')


doc.add_heading('6. İlerleme ve tekrar oynama sistemleri', level=1)
doc.add_heading('6.1 Dört zorluk modu', level=2)
mode_rows = [
    ('Başlangıç', '2 → 3', '1 → 2', '42 sn', '1×', 'Rüzgâr yok; priority ve ileri komut yok; sürekli yardım.'),
    ('Normal', '3,5 → 8', '3 → 8', '18 sn', '2×', 'Dengeli trafik; priority açık; ileri komutlar kapalı.'),
    ('İleri', '8 → 15', '3 → 12', '11 sn', '2×', 'Tam rüzgâr; published procedures, DCT/HOLD, olaylar.'),
    ('Uzman', '12 → 30', '3 → 20', '6,5 sn', '2×', 'Yüksek baskı; sert hava; pist kontrolü ve tam operasyon.'),
]
add_table(doc, ['Mod', 'Skill', 'Uçak', 'Min spawn', 'Hız', 'Karakter'], mode_rows, [1200, 1050, 950, 1050, 650, 4460], font_size=8.6,
          alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])

doc.add_heading('6.2 Vardiya hedefleri', level=2)
goal_rows = [
    ('Başlangıç', '1', '0', '0'),
    ('Normal', '3', '1', '≤1'),
    ('İleri', '6', '2', '≤1'),
    ('Uzman', '10', '4', '0'),
]
add_table(doc, ['Mod', 'İniş', 'Handoff', 'Maks. separation loss'], goal_rows, [2200, 1500, 1700, 3960], font_size=9.2,
          alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER])

doc.add_heading('6.3 Mini Metro tarzı havalimanı açma', level=2)
add_body(doc, 'İstanbul doğrudan açıktır. Her sonraki havalimanı, kariyer sırasındaki bir önceki havalimanında belirli en iyi skora ulaşınca açılır. Kapı formülü: 250 + floor((index−1)/6) ×100.')
unlock_rows = [
    ('Sıra 2–7', 'Önceki havalimanında 250'),
    ('Sıra 8–13', '350'),
    ('Sıra 14–19', '450'),
    ('Sıra 20–25', '550'),
    ('Sıra 26–31', '650'),
    ('Sıra 32–37', '750'),
    ('Sıra 38–43', '850'),
    ('Sıra 44–49', '950'),
    ('Sıra 50', '1.050'),
]
add_table(doc, ['Açılacak sıra', 'Gerekli önceki meydan skoru'], unlock_rows, [2600, 6760], font_size=9.2)
add_body(doc, 'Normal mod “beginner-complete”; İleri mod “normal-complete + procedure-pilot”; Uzman mod “advanced-complete + clean-shift” başarımlarıyla açılır. Wake Advisor, Flow Management ve Master Flow gibi operasyon araçları da başarımlarla kapılanır.')

doc.add_heading('6.4 Flight Academy: 10 ders', level=2)
academy_rows = [
    ('01', 'Radarı Oku', '2 dk', 'Radar hedefini seç; etiket bilgisini oku.'),
    ('02', 'Heading Ver', '2 dk', 'Seçili uçağın heading değerini değiştir.'),
    ('03', 'İrtifayı Yönet', '2 dk', 'Yeni irtifa ver ve kademeli alçalt.'),
    ('04', 'Hızı Ayarla', '2 dk', 'Yaklaşma sırasını hızla düzenle.'),
    ('05', 'ILS Yaklaşması', '3 dk', 'Aktif piste ILS’i silahlandır.'),
    ('06', 'Localizer Yakala', '3 dk', 'Uygun açı ve alttan glideslope yaklaşımı.'),
    ('07', 'İlk İniş', '3 dk', 'İlk güvenli inişi ve otomatik tower handoff’u tamamla.'),
    ('08', 'Kalkışı Devret', '2 dk', 'Departure handoff onayı ver.'),
    ('09', 'Direct ve Hold', '3 dk', 'Aynı uçakta DCT ve HOLD uygula.'),
    ('10', 'Güvenli Sıralama', '3 dk', 'İki gelişe yaklaşma ver; ayırma ve wake’i koru.'),
]
add_table(doc, ['#', 'Ders', 'Süre', 'Amaç'], academy_rows, [500, 2000, 900, 5960], font_size=8.9,
          alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])

doc.add_heading('6.5 Istanbul Control — First Watch', level=2)
career_rows = [
    ('01', 'İlk Temas', 'Başlangıç · north-parallel', '1 iniş, 150 puan, 0 kayıp', 'Supervisor mesajı'),
    ('02', 'Paralel Hatlar', 'Normal · north-parallel', '3 iniş, 1 handoff, 450 puan', 'Demand pulse'),
    ('03', 'Sis Hattı', 'Normal · north-lowvis', '3 iniş, 1 handoff, 440 puan', 'Tek pist / düşük görüş'),
    ('04', 'Öncelik Bir', 'İleri · north-parallel', '4 iniş, 1 handoff, 600 puan', 'Medical priority'),
    ('05', 'Pist Dönüşü', 'İleri · north-parallel', '5 iniş, 2 handoff, 720 puan', 'Flow change → north-single'),
    ('06', 'Gece Dalgası', 'İleri · south-triple', '6 iniş, 2 handoff, 850 puan', 'Heavy long-haul demand pulse'),
    ('07', 'Baş Kontrolör', 'Uzman · north-parallel', '7 iniş, 3 handoff, 1.100 puan', 'Minimum fuel + final demand pulse'),
]
add_table(doc, ['#', 'Bölüm', 'Mod / flow', 'Hedef', 'Ana olay'], career_rows, [450, 1550, 2100, 2850, 2410], font_size=8.4)
add_body(doc, 'Bölümler sırayla açılır. Debrief üç sonuç üretir: Distinction / Üstün Başarı, Qualified / Yeterli ve Repeat / Tekrar Gerekli. Olay kimlikleri eventTimeline içinde idempotent tutulur; kayıt geri yüklenince aynı olay tekrar tetiklenmez.')

doc.add_heading('6.6 Daily Radar, streak, logbook ve paylaşım', level=2)
for text in [
    'UTC gününe göre IST, LHR, LAX, JFK veya ATL; geçerli flow ve seed deterministik seçilir.',
    'Daily Radar Normal moddadır: 3 iniş, 1 handoff, 0 separation loss ve 350/375/400 puan hedefi.',
    'Aynı tarih bir kez tamamlanır; güncel ve en iyi streak yerelde hesaplanır.',
    'Logbook en yeni 30 vardiyayı saklar; skor, iniş, handoff, peak skill, grade ve emniyet metriklerini içerir.',
    'Web Share destekleniyorsa native paylaşım sayfası; değilse metin panoya kopyalanır.',
    'Şu an hiçbir veri sunucuya yüklenmez.',
]:
    add_bullet(doc, text, num_ids)


doc.add_heading('7. Havalimanı, pist ve prosedür verisi', level=1)
doc.add_heading('7.1 Kaynak ve doğruluk politikası', level=2)
for text in [
    'Yolcu sırası ACI 2025 dünya trafik sıralamasından gelir; İstanbul ürün kararıyla kariyer sırası 1’e taşınmıştır.',
    'Pist eşik koordinatları, yönleri, uzunlukları ve yükseklikler OurAirports public-domain snapshot commit be07e33e6cc10087f57064f2bb3fccfcd39f5801 kaynaklıdır.',
    'Her fiziksel pist çifti tek çizgi olarak modellenir; reciprocal flow aynı çizgiyi ters yönde kullanır.',
    'Her havalimanında normal, reverse ve reduced/low-vis flow bulunur; flagship paketlerde dördüncü araştırılmış flow vardır.',
    'Şehir, su, dağ ve arazi sınıfları taktik yön bağlamıdır; gerçek poligon veya survey-grade terrain değildir.',
    'Her pist sonu için ILS-aligned oyun final koridoru vardır; bu, yayımlanmış yaklaşmanın birebir kopyası değildir.',
    'Chart-derived rotalar kompakt 40–42 NM oyun sektörüne projekte edilir; seyrüsefer verisi değildir.',
]:
    add_bullet(doc, text, num_ids)

doc.add_heading('7.2 İlk beş flagship operations pack', level=2)
flagship_rows = [
    ('IST / LTFM', 'Triple-independent runway bank', 'Black Sea wind shift; tek piste düşüş ve üçlü akışa dönüş', '4 geliş + 2 kalkış / 6; heavy cadence 4'),
    ('LHR / EGLL', 'BNN/LAM/BIG/OCK dört stack merge', 'Runway alternation bütün geliş sırasını taşır', '4 geliş + 1 kalkış / 5; heavy cadence 3'),
    ('LAX / KLAX', 'North/south complex balancing', 'North complex bakım; south complex’e sıkışma', '3 geliş + 2 kalkış / 5; heavy cadence 4'),
    ('JFK / KJFK', 'Intersecting 22/31 demand', 'Low visibility tek geliş hattı', '3 geliş + 2 kalkış / 5; heavy cadence 3'),
    ('ATL / KATL', 'Five-parallel throughput', 'PRM-style capacity compression/recovery', 'Alternating traffic; heavy cadence 6'),
]
add_table(doc, ['Meydan', 'Kimlik', 'Olay', 'Trafik karakteri'], flagship_rows, [1300, 2350, 2850, 2860], font_size=8.5)

doc.add_heading('7.3 Yayımlanmış prosedür kapsamı', level=2)
add_callout(doc, 'KAPSAM', '29/50 havalimanında yayımlanmış runtime route vardır (%58). Kalan 21 havalimanı açıkça generated vector route kullanır; gerçek prosedür gibi etiketlenmez.')
remaining = [(iata, icao) for _, iata, icao, _, _ in AIRPORTS if icao not in PUBLISHED_ICAO]
remaining_text = ', '.join(f'{iata}/{icao}' for iata, icao in remaining)
add_body(doc, f'Yayımlanmış runtime rotası bulunmayan 21 meydan: {remaining_text}.')
add_body(doc, 'ZUTF ve WIII için resmî CAAC/Indonesia PIA içerikleri authenticated erişim gerektirdiğinden üçüncü taraf veya eski chart kullanılmamıştır. Bu güvenli veri ilkesi diğer meydanlarda da korunmalıdır.')

doc.add_heading('7.4 FAA CIFP yenileme hattı', level=2)
for text in [
    'FAA’nın 28 günlük cycle ile yayımladığı yaklaşık 53 MB, 132 kolonlu ARINC 424-18 kaynak dosyası repoya commit edilmez.',
    'npm run import:cifp -- /path/to/FAACIFP18 komutu seçili ABD paketini yeniden üretir.',
    'Importer record width doğrular, airport/fix koordinatlarını çözer, feeder transition ve common route’u birleştirir, altitude/speed kısıtlarını decode eder ve runway compatibility çıkarır.',
    'Deterministik çıktı src/engine/generated/faaCifpProcedures.ts altında commit edilir; cycle değişikliği normal diff olarak incelenebilir.',
]:
    add_bullet(doc, text, num_ids)


doc.add_heading('8. Teknik mimari', level=1)
tech_rows = [
    ('Frontend', 'React + TypeScript + Vite'),
    ('Radar', 'Canvas 2D; responsive viewport ve pointer/touch interaction'),
    ('Motor', 'UI’dan ayrılmış saf TypeScript simülasyonu'),
    ('Test', 'Vitest; Node environment'),
    ('Storage', 'Browser localStorage; session version 4'),
    ('Offline', 'PWA manifest + service worker cache airspace-control-core-v9-first-watch'),
    ('Native hedef', 'Capacitor; appId com.bugrabilim.airspacecontrol; webDir dist'),
    ('Hosting', 'Vercel; Vite framework; Node 24.x build environment'),
    ('Backend', 'Yok; Supabase deferred'),
]
add_table(doc, ['Katman', 'Teknoloji / karar'], tech_rows, [1700, 7660], font_size=9.3)

doc.add_heading('8.1 Motor modülleri', level=2)
module_rows = [
    ('academy.ts', '10 derslik deterministik eğitim board’u ve değerlendirme.'),
    ('aircraftData.ts', 'Jet/heavy/light performans ve wake kategori eşleme.'),
    ('aircraftDynamics.ts', 'Bank, roll, speed, wind, track, vertical motion ve predictor.'),
    ('airportCatalog.ts', '50 havalimanı, pistler, trafik, terrain/city/water/mountain bağlamı.'),
    ('airportOperations.ts', 'IST/LHR/LAX/JFK/ATL flagship operations ve kaynak manifestleri.'),
    ('publishedProcedureCatalog.ts', 'FAA ve uluslararası chart-derived pack birleşimi.'),
    ('approach.ts', 'Localizer/glideslope/tower/go-around geometrisi ve durum makinesi.'),
    ('arrivalAdvisor.ts', 'Pist bazlı yaklaşma sırası ve radar karar önerisi.'),
    ('commands.ts', 'Parser, compact alias, callsign resolution, multi-command ve apply.'),
    ('controllerCoach.ts', 'Aktif trafiğe göre bağlamsal öneri.'),
    ('difficulty.ts', '4 modun trafik, rüzgâr, skill ve komut sınırları.'),
    ('engagement.ts', 'Daily Radar, streak, 30 kayıt logbook ve paylaşım.'),
    ('careerSeason.ts', 'First Watch bölümleri, event trigger/effect ve sonuçlar.'),
    ('navigation.ts', 'Route, DCT ve HOLD guidance.'),
    ('progression.ts', 'Skor, debrief, 52 başarım ve unlock haritası.'),
    ('radio.ts', 'NATO callsign, TTS metni ve readback queue gecikmesi.'),
    ('scenario.ts', 'Catalog’u oynanabilir RadarWorld ve başlangıç state’ine dönüştürür.'),
    ('separation.ts / wake.ts', 'Ayırma, CPA, paralel/diverging istisnalar ve A–F wake matrisi.'),
    ('session.ts', 'Version 1–4 restore/migration ve session serialization.'),
    ('simulation.ts', '0,05 sn orchestration, eventler, pist, trafik, scoring ve state step.'),
    ('skill.ts', 'Canlı performans puanı ve trafik profili.'),
    ('trafficDirector.ts', 'Seed’li arrival/departure planı, runway load ve published route seçimi.'),
    ('weather.ts', 'Flow weather, runway wind components ve approach toleransı.'),
]
add_table(doc, ['Dosya', 'Sorumluluk'], module_rows, [2300, 7060], font_size=8.7)

doc.add_heading('8.2 UI modülleri', level=2)
ui_rows = [
    ('LandingPage', 'English-first marketing, career map, daily, academy, airports, footer.'),
    ('App', 'Ürün state’i, persistence, audio, routing ve ana orchestration.'),
    ('RadarScope', 'Canvas radar, pan/zoom/lock, drag labels, predictor, measurement.'),
    ('CommandPanel', 'Desktop command deck ve mobile bottom sheet.'),
    ('FlightStripList', 'Uçuş listesi, conflict/priority göstergeleri.'),
    ('MissionPanel', 'Görev, score, unlock, flow/procedure ve coach detayları.'),
    ('AcademyPanel', 'Ders ilerleme ve eylem CTA’ları.'),
    ('CareerSeasonPanel', '7 bölümün kilit/açık/sonuç görünümü.'),
    ('DailyChallengePanel', 'Günlük görev, streak ve logbook.'),
    ('DebriefPanel', 'Grade, başarımlar, share ve devam seçenekleri.'),
]
add_table(doc, ['Bileşen', 'Sorumluluk'], ui_rows, [2200, 7160], font_size=9.0)

doc.add_heading('8.3 Persistence', level=2)
storage_rows = [
    ('airspace-control-career-v1', 'Best score/landings, badges, airport best scores, Daily dates, logbook, career outcomes.'),
    ('airspace-control-session-v1', 'Aktif vardiya; her 3 sn ve beforeunload sırasında save.'),
    ('airspace-control-academy-v1', 'Tamamlanan Academy lesson id’leri.'),
]
add_table(doc, ['localStorage anahtarı', 'İçerik'], storage_rows, [3100, 6260], font_size=9.2)
add_body(doc, 'Session version 4’tür; restore versions 1–3 için migration içerir. Restore edilen vardiya güvenlik amacıyla paused başlar. Private mode veya quota sorunu storage’ı engellerse uygulama sessizce çalışmaya devam eder fakat kalıcılık garanti edilmez.')

doc.add_heading('8.4 Web, PWA ve native paketleme', level=2)
for text in [
    'npm install → npm run dev ile yerel web geliştirme.',
    'npm run typecheck, npm test ve npm run build yayın öncesi zorunlu kapılardır.',
    'PWA manifest standalone/minimal-ui, theme/background renkleri, maskable SVG icon ve any orientation içerir.',
    'Service worker navigation için network-first/offline fallback; assetler için cache-first + background refresh davranışı kullanır.',
    'Capacitor config hazırdır; fakat @capacitor/core/cli/android/ios paketleri mevcut package.json dependencies içine eklenmemiştir.',
    'Google Play için Android Studio imzalı .aab; App Store için macOS/Xcode imzalı .ipa, developer hesapları, sertifikalar ve gizlilik formları gerekir.',
]:
    add_bullet(doc, text, num_ids)


doc.add_heading('9. FlatOut ATC ve Endless ATC araştırma sonucu', level=1)
doc.add_heading('9.1 FlatOut ATC’den öğrenilen ürün sistemi', level=2)
for text in [
    'Başarı tek başına radar çiziminden değil; kısa Academy, gerçek havalimanı kimliği, farklı görüş modları, kariyer, günlük içerik, logbook, olaylar ve live traffic ekosisteminden gelir.',
    'Yeni oyuncu deneyimi uzun manuel yerine yaklaşık üç dakikalık tek beceri dersleriyle kurulmalıdır.',
    'Havalimanı sayısı tek başına yeterli değildir; her meydan farklı bir operasyon bulmacası olmalıdır.',
    'Mobilde komut verme düşünce hızına yaklaşmalı; ekranın çoğu radar kalmalıdır.',
    'Vardiya sonrası debrief ve kalıcı kayıt, oyuncunun yaptığı işe anlam kazandırır.',
    '3D radar/tower, live ADS-B, ground, helikopter/SAR ve oceanic operations ileri ürün katmanlarıdır; çekirdek 2D kalitesinden sonra gelmelidir.',
]:
    add_bullet(doc, text, num_ids)

doc.add_heading('9.2 Endless ATC’den öğrenilen çekirdek', level=2)
for text in [
    'Adaptif skill–traffic geri beslemesi oyunun sonsuz döngüsünün merkezidir.',
    'Heading/irtifa/hız kararları bank, roll, speed ve wind fiziğiyle gerçek sonuç üretmelidir.',
    'ILS tek düğme değil; localizer ve glideslope capture şartları olan bir durum makinesi olmalıdır.',
    'Ayırma, wake, pist kapasitesi ve hızlı komut girişi yoğun trafiği anlamlı kılar.',
    'Radar pan/zoom, movable label, predictor, track history ve measurement araçları okunabilirliği korur.',
]:
    add_bullet(doc, text, num_ids)

doc.add_heading('9.3 Bugünkü işlevsel eşleşme', level=2)
parity_rows = [
    ('2D radar / vektör / fizik', 'Büyük ölçüde uygulanmış', 'Mobil gerçek cihaz acceptance ve uzun vardiya soak testi gerekli.'),
    ('ILS / separation / wake', 'Uygulanmış', 'Edge case ve havalimanı bazlı yaklaşma doğrulaması genişletilmeli.'),
    ('Eğitim', '10 Academy dersi uygulanmış', 'Gerçek kullanıcı onboarding testi gerekli.'),
    ('Kariyer / daily / logbook', 'Local-first uygulanmış', 'Bulut sync ve leaderboard yok.'),
    ('Havalimanı operasyon kimliği', '5 flagship güçlü; diğerleri değişken', '21 meydan generated vector; tüm meydanlarda derin operasyon paketi yok.'),
    ('Published procedures', '29/50', 'SID, yaklaşma ve missed approach kapsamı çoğu meydanda eksik.'),
    ('Olaylar / aciller', 'Temel deterministik set var', 'Dynamic thunderstorms, volcanic ash, 7600, engine failure yok.'),
    ('3D / tower / ground', 'Yok', 'İleri faz.'),
    ('Live ADS-B', 'Yok', 'Veri lisansı, maliyet ve fallback tasarımı gerektirir; şart değil.'),
    ('Store dağıtımı', 'Capacitor config var', 'Native dependencies, icons/splash, signing ve mağaza hazırlığı yapılmadı.'),
]
add_table(doc, ['Sistem', 'Durum', 'Açık'], parity_rows, [2200, 2350, 4810], font_size=8.7)

doc.add_heading('9.4 Hukuk ve ürün kimliği sınırı', level=2)
add_callout(doc, 'ÖNEMLİ', 'İsim ve renk değişikliği tek başına hukuki veya mağaza incelemesi açısından yeterli değildir. FlatOut ATC’nin kodu, binary’si, ekran düzeni, metinleri, hikâyesi, ikonları, sesleri veya görselleri kopyalanmamalıdır.', tone='warning')
for text in [
    'Genel ATC yöntemleri, heading/altitude/speed komutları, vardiya, günlük görev ve leaderboard gibi fikirler bağımsız uygulanabilir.',
    'Apple App Review Guidelines 4.1 Copycats ve 5.2 Intellectual Property mağaza öncesi tekrar kontrol edilmelidir.',
    'Açık yayımlanmış havacılık bilgisi otomatik olarak sınırsız yeniden dağıtım hakkı anlamına gelmez; her veri kaynağında lisans/provenans saklanmalıdır.',
    'GAME ONLY ve NOT FOR NAVIGATION sınırı ürün içinde görünür kalmalıdır.',
]:
    add_bullet(doc, text, num_ids)


doc.add_heading('10. Doğrulama, açıklar ve riskler', level=1)
doc.add_heading('10.1 26 Ağustos 2026 doğrulaması', level=2)
verification_rows = [
    ('GitHub main', 'b95ec92 doğrulandı'),
    ('Vitest', '20 test dosyası · 137/137 test geçti'),
    ('TypeScript', 'tsc -b --pretty false temiz'),
    ('Vite build', '55 module transformed · build başarılı'),
    ('JS bundle', '520,56 kB minified · 143,64 kB gzip'),
    ('CSS bundle', '83,89 kB · 15,55 kB gzip'),
    ('Vercel', 'Production READY'),
    ('Runtime errors', 'Son 24 saatte hata kümesi yok'),
]
add_table(doc, ['Kontrol', 'Sonuç'], verification_rows, [2300, 7060], font_size=9.3)

doc.add_heading('10.2 Bilinen teknik ve ürün açıkları', level=2)
risks = [
    ('P0', 'Fiziksel mobil QA eksik', 'iPhone Safari/Chrome ve gerçek Android cihazda portrait/landscape, keyboard, safe area, audio, scroll ve 30 dakikalık vardiya acceptance yapılmadı.'),
    ('P0', 'CSS katmanlanması', 'styles.css yaklaşık 1.705 satır ve birden fazla tarihsel responsive pass içeriyor; cascade çakışması riski yüksek.'),
    ('P0', 'Görsel ürün uyumu', 'Kullanıcı mevcut tasarımın FlatOut ile yeterince ilişkili görünmediğini belirtti; yeniden tasarım kopya olmadan ürün kalitesi ve radar ergonomisine odaklanmalı.'),
    ('P1', 'Procedure coverage', '21 havalimanında gerçek runtime route yok; 29 meydanın çoğunda arrival ağırlıklı, SID/approach/missed approach eksik.'),
    ('P1', 'Terrain / city doğruluğu', 'Taktik bearings var; gerçek kent, su, dağ ve restricted area poligonları yok.'),
    ('P1', 'Vector vs route kararı', 'Published route auto-follow davranışı oyuncu iş yükünü azaltabilir; hedef loop açısından tekrar ürün kararı gerekir.'),
    ('P1', 'Bundle', '520 kB JS chunk Vite uyarısı veriyor; landing/game, research-heavy data ve paneller code-split edilmeli.'),
    ('P1', 'Ses', 'Varsayılan açık olsa da browser gesture şartı ve device voice availability nedeniyle ilk açılışta sessizlik yaşanabilir.'),
    ('P2', 'Cloud yok', 'Cihazlar arası sync, hesap, leaderboard ve telemetry yok.'),
    ('P2', 'Native paket yok', 'Capacitor config dışında Android/iOS proje klasörleri ve imzalı çıktılar yok.'),
    ('P2', 'Gelişmiş operasyonlar yok', 'Ground/taxi, 3D tower, live ADS-B, SAR/heli/oceanic ve gelişmiş weather yok.'),
]
add_table(doc, ['Öncelik', 'Açık', 'Etkisi'], risks, [850, 2150, 6360], font_size=8.7,
          alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])


doc.add_heading('11. Öncelikli uygulama planı', level=1)
add_callout(doc, 'ÇALIŞMA KURALI', 'Her faz temiz origin/main üzerinde yapılmalı; ilgili testler ve görsel doğrulama tamamlanmadan sonraki faza geçilmemeli. GitHub/Vercel gönderimi kullanıcı onayı gerektirir.')

doc.add_heading('P0 — Kaynak ve mobil ürün stabilizasyonu', level=2)
reset_number_sequence(doc, num_ids)
for text in [
    'Eski dirty yerel dalı güvenli biçimde karşılaştır; çalışma için temiz origin/main kopyası kullan.',
    'Gerçek iPhone ve Android test matrisi oluştur: 320–430 px portrait, küçük/uzun ekran, landscape, virtual keyboard, safe areas ve browser chrome.',
    'Radarın minimum görünür alanını, aircraft label boyutunu ve 44×44 px dokunma hedeflerini kabul kriteri yap.',
    'Mobile command sheet aç/kapat, selection, scroll ve keyboard akışını video/screenshot ile doğrula.',
    'Audio unlock için ilk Play/Continue etkileşiminde açık geri bildirim ver; TTS ve cue fallback’ini test et.',
    'styles.css responsive katmanlarını tek canonical mobile/desktop sisteme indir; görsel regression ekle.',
]:
    add_numbered(doc, text, num_ids)

doc.add_heading('P1 — Oynanış ve görsel kalite eşleşmesi', level=2)
reset_number_sequence(doc, num_ids)
for text in [
    'FlatOut’ın görünümünü kopyalamadan, bilgi hiyerarşisi ve dokunma hızını referans alan yeni radar HUD tasarımını tamamla.',
    'Vector-first ile procedure-assisted davranış için bir ürün kararı ver; modlara göre ayrıştırmayı değerlendir.',
    'Departure release/pist aralığı gibi oyuncu kararlarını genişlet; kalkışın yalnızca handoff hedefi olmasını düzelt.',
    '30–60 dakikalık soak testlerde spawn, save/restore, events, traffic metering, audio queue ve memory davranışını doğrula.',
    'Bundle’ı landing, game, procedure data ve secondary panels olarak code-split et.',
]:
    add_numbered(doc, text, num_ids)

doc.add_heading('P2 — Havalimanı ve prosedür tamamlama', level=2)
reset_number_sequence(doc, num_ids)
for text in [
    'Kalan 21 meydan için yalnızca current authoritative source erişilebiliyorsa STAR/SID/approach paketi ekle.',
    'Mevcut 29 meydanda arrival-only kapsamı SID, approach, missed approach ve runway-flow compatibility ile genişlet.',
    'Her pack için authority, cycle/effective date, accessed date, license/distribution note ve automated geometry/constraint tests ekle.',
    'Terrain/city/water/mountain katmanını lisanslı veya açık veriyle gerçek poligonlara yükselt; gameplay projection ile açıkça ayır.',
]:
    add_numbered(doc, text, num_ids)

doc.add_heading('P3 — Ürün derinliği', level=2)
reset_number_sequence(doc, num_ids)
for text in [
    'Dynamic thunderstorms/wind shift, 7600, engine failure ve additional emergency scenario seti.',
    'Daha fazla özgün kariyer sezonu ve havalimanına özel görevler.',
    'İsteğe bağlı ground/taxi ve tower alt modu.',
    'Analytics olmadan önce gizlilik, event taxonomy ve açık kullanıcı rızası tasarımı.',
]:
    add_numbered(doc, text, num_ids)

doc.add_heading('P4 — Cloud ve mağaza', level=2)
reset_number_sequence(doc, num_ids)
for text in [
    'Supabase veya alternatif backend: passwordless auth, cloud save, leaderboard ve migration planı.',
    'Capacitor native dependencies, Android/iOS projeleri, adaptive icons, splash, offline/audio/device QA.',
    'Privacy policy, terms, data safety/App Privacy formları, age rating, screenshots ve store metadata.',
    'Android signed AAB ve iOS archive/IPA; TestFlight ve Play Internal Testing kabul turu.',
]:
    add_numbered(doc, text, num_ids)

doc.add_heading('11.1 P0 tamamlanma ölçütü', level=2)
for text in [
    'atc-tr.vercel.app landing ve atc-tr-play.vercel.app game route 200/READY.',
    'iPhone ve Android portrait/landscape ekranlarında radar okunabilir; yatay taşma yok.',
    'Uçak seçme ve temel heading/altitude/speed/ILS komutu üç dokunuş içinde tamamlanabilir.',
    'Ses ilk oyun etkileşimi sonrasında duyulur veya kullanıcıya açık cihaz/browser nedeni gösterilir.',
    '137+ test, typecheck ve build temiz; yeni görsel/mobile testler eklenmiş.',
    '30 dakikalık vardiyada crash, trafik kilitlenmesi, görünmez panel veya save/restore bozulması yok.',
]:
    add_bullet(doc, text, num_ids)


doc.add_heading('12. Yeni sohbet için dosya ve komut rehberi', level=1)
path_rows = [
    ('README.md', 'Ürün özeti, komutlar, kurulum ve native paketleme notu.'),
    ('docs/AIRPORT_DATA.md', '50 havalimanı sırası, pist veri politikası ve kaynaklar.'),
    ('docs/PUBLISHED_PROCEDURES.md', '29 meydan route kapsamı ve FAA import.'),
    ('docs/FLAGSHIP_AIRPORT_PACKS.md', 'IST/LHR/LAX/JFK/ATL araştırma ve operasyon karakteri.'),
    ('docs/FIRST_WATCH_SEASON.md', '7 bölümlük kariyer sözleşmesi.'),
    ('docs/RETENTION_SYSTEMS.md', 'Daily Radar, streak, logbook ve share.'),
    ('docs/flatout-atc-research-tr.md', 'FlatOut derin araştırma, parity ve hukuk sınırı.'),
    ('docs/endless-atc-parity-report-tr.md', 'Endless çekirdek oynanış analizi.'),
    ('src/engine', 'Platform-neutral oyun motoru.'),
    ('src/ui', 'React UI, Canvas radar ve responsive CSS.'),
]
add_table(doc, ['Yol', 'Amaç'], path_rows, [3300, 6060], font_size=9.0)

doc.add_heading('Standart doğrulama komutları', level=2)
for command in ['npm install', 'npm run dev', 'npm run typecheck', 'npm test', 'npm run build', 'npm run import:cifp -- /path/to/FAACIFP18']:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=4)
    set_run_font(p.add_run(command), name='Consolas', size=9.5, color=INK, bold=True)

doc.add_heading('Yeni sohbet ilk kontrol listesi', level=2)
reset_number_sequence(doc, num_ids)
for text in [
    'git ls-remote origin refs/heads/main ile SHA’yı doğrula.',
    'Vercel production deployment, domain aliasları ve son runtime error durumunu kontrol et.',
    'Temiz origin/main snapshot üzerinde test/typecheck/build çalıştır.',
    'Bu belgedeki 50/29/52/137 sayıları yeni commit varsa yeniden hesapla.',
    'Dirty yerel dalı koru; hiçbir dosyayı silme veya resetleme.',
    'P0 işini uygula, görsel ve otomatik doğrula; push/deploy için onay iste.',
]:
    add_numbered(doc, text, num_ids)


doc.add_page_break()
doc.add_heading('Ek A — 50 havalimanı kariyer kataloğu', level=1)
airport_rows = []
for rank, iata, icao, passengers, runways in AIRPORTS:
    coverage = 'Yayımlanmış route' if icao in PUBLISHED_ICAO else 'Generated vector'
    airport_rows.append((rank, iata, icao, passengers, runways, coverage))
add_table(doc, ['#', 'IATA', 'ICAO', '2025 yolcu', 'Pist çifti', 'Runtime route'], airport_rows,
          [500, 850, 950, 2100, 1200, 3760], font_size=8.0,
          alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])

doc.add_page_break()
doc.add_heading('Ek B — 29 yayımlanmış runtime prosedür paketi', level=1)
add_table(doc, ['ICAO', 'Runtime prosedür', 'Authority / cycle'], PROCEDURES,
          [1100, 5180, 3080], font_size=7.9)

doc.add_heading('Ek C — 52 başarım', level=1)
achievement_rows = [(idx, label, description) for idx, (_, label, description) in enumerate(ACHIEVEMENTS, start=1)]
add_table(doc, ['#', 'Başarım', 'Koşul'], achievement_rows,
          [500, 2600, 6260], font_size=8.3,
          alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])

doc.add_page_break()
doc.add_heading('Ek D — Kaynaklar ve bağlantılar', level=1)
sources = [
    ('Canlı ürün', 'https://atc-tr.vercel.app'),
    ('GitHub deposu', 'https://github.com/bugrabilim/atc'),
    ('Son commit', 'https://github.com/bugrabilim/atc/commit/b95ec92ab64e6098f66f13b3ff72c43810aafe81'),
    ('ACI busiest airports', 'https://aci.aero/resources/busiest-airports-in-the-world/'),
    ('OurAirports data', 'https://ourairports.com/data/'),
    ('FAA d-TPP', 'https://www.faa.gov/air_traffic/flight_info/aeronav/digital_products/dtpp/'),
    ('DHMI İstanbul Airport', 'https://www.dhmi.gov.tr/sayfalar/havalimani/istanbul/GenelBilgiler.aspx'),
    ('FlatOut ATC', 'https://flatoutatc.com/'),
    ('Apple App Review Guidelines', 'https://developer.apple.com/app-store/review/guidelines/'),
]
for label, url in sources:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=5)
    set_run_font(p.add_run(f'{label}: '), bold=True, color=INK)
    add_hyperlink(p, 'Kaynağı aç', url)

add_callout(doc, 'SON NOT', 'Bu master dosya güncel bir devir anlık görüntüsüdür. Yeni sohbette gerçek kaynak GitHub main ve canlı Vercel durumudur. Yeni kod yazmadan önce bu iki kaynak yeniden doğrulanmalıdır.')


doc.save(OUTPUT)
print(OUTPUT)
